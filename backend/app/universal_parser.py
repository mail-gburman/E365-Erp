"""
Universal File Parser — reads any file format and extracts tabular data.
Supports: .xlsx, .xls, .csv, .tsv, .pdf, .docx, .jpg, .jpeg, .png
"""
import csv
import re
from io import BytesIO, StringIO
from difflib import SequenceMatcher


# ── Known field schemas per entity type ──
ENTITY_SCHEMAS = {
    "crew": {
        "fields": {
            "full_name":      {"label": "Full Name",       "required": True,  "aliases": ["name", "employee name", "crew name", "staff name", "manpower name", "person"]},
            "employee_code":  {"label": "Employee Code",   "required": False, "aliases": ["code", "emp code", "staff code", "emp id", "id"]},
            "role":           {"label": "Role/Designation","required": False, "aliases": ["designation", "position", "title", "dept"]},
            "manpower_type":  {"label": "Manpower Type",   "required": False, "aliases": ["type", "employment type", "emp type", "category"]},
            "phone":          {"label": "Phone/Mobile",    "required": False, "aliases": ["contact number", "contact no", "mobile", "mob", "cell", "tel", "telephone", "ph no", "phone no", "phone number", "contact"]},
            "home_station":   {"label": "Home Station",    "required": False, "aliases": ["station", "base", "city", "location", "home city"]},
            "address":        {"label": "Address",         "required": False, "aliases": ["addr", "home address", "residential address"]},
            "aadhar_number":  {"label": "Aadhar Number",   "required": False, "aliases": ["aadhar", "aadhaar", "aadhar no", "aadhaar no", "aadhar id", "uid"]},
            "id_proof_type":  {"label": "ID Proof Type",   "required": False, "aliases": ["id type", "proof type", "document type"]},
            "id_proof_number":{"label": "ID Proof Number", "required": False, "aliases": ["id number", "proof number", "document number", "id no"]},
        },
    },
    "inventory": {
        "fields": {
            "name":              {"label": "Equipment Name","required": True,  "aliases": ["equipment", "item", "item name", "device", "description", "equipment name", "product"]},
            "asset_code":        {"label": "Asset Code",    "required": False, "aliases": ["unique number", "unique no", "asset id", "code", "unique", "asset"]},
            "serial_number":     {"label": "Serial Number", "required": False, "aliases": ["serial", "serial no", "sl no", "sr no", "s/n", "sn"]},
            "category":          {"label": "Category",      "required": False, "aliases": ["cat", "equipment category", "type", "group"]},
            "item_type":         {"label": "Item Type",     "required": False, "aliases": ["device type", "classification"]},
            "brand":             {"label": "Brand/Make",    "required": False, "aliases": ["make", "manufacturer", "mfg"]},
            "model_no":          {"label": "Model No.",     "required": False, "aliases": ["model", "model number"]},
            "quantity":          {"label": "Quantity",      "required": False, "aliases": ["qty", "qty nos", "nos", "count", "units"]},
            "location_text":     {"label": "Location",      "required": False, "aliases": ["location", "rack", "shelf", "bay", "stored at"]},
            "warranty_expiry":   {"label": "Warranty Expiry","required": False,"aliases": ["warranty", "warranty date", "expiry", "warranty end"]},
            "purchase_date":     {"label": "Purchase Date", "required": False, "aliases": ["purchase", "date of purchase", "bought on"]},
            "notes":             {"label": "Remarks",       "required": False, "aliases": ["remarks", "remark", "comments", "comment", "note"]},
        },
    },
    "clients": {
        "fields": {
            "name":            {"label": "Client Name",    "required": True,  "aliases": ["client", "company", "company name", "party name", "firm"]},
            "client_code":     {"label": "Client Code",    "required": False, "aliases": ["code", "id"]},
            "industry_type":   {"label": "Industry Type",  "required": False, "aliases": ["industry", "sector", "type"]},
            "billing_address": {"label": "Billing Address","required": False, "aliases": ["address", "addr", "office address"]},
            "gst_number":      {"label": "GST Number",    "required": False, "aliases": ["gst", "gstin", "gst no"]},
            "notes":           {"label": "Notes",          "required": False, "aliases": ["remarks", "note", "comments"]},
        },
    },
    "vendors": {
        "fields": {
            "name":            {"label": "Vendor Name",    "required": True,  "aliases": ["vendor", "company", "company name", "supplier", "firm"]},
            "vendor_code":     {"label": "Vendor Code",    "required": False, "aliases": ["code", "id"]},
            "vendor_type":     {"label": "Vendor Type",    "required": False, "aliases": ["type", "category"]},
            "city":            {"label": "City",           "required": False, "aliases": ["location", "place"]},
            "contact_person":  {"label": "Contact Person", "required": False, "aliases": ["contact name", "contact", "person"]},
            "phone":           {"label": "Phone",          "required": False, "aliases": ["contact number", "mobile", "tel", "phone no"]},
            "email":           {"label": "Email",          "required": False, "aliases": ["email id", "mail", "e-mail"]},
            "gst_number":      {"label": "GST Number",     "required": False, "aliases": ["gst", "gstin", "gst no"]},
            "notes":           {"label": "Notes",          "required": False, "aliases": ["remarks", "note", "comments"]},
        },
    },
    "warehouses": {
        "fields": {
            "name":          {"label": "Warehouse Name", "required": True,  "aliases": ["warehouse", "store", "godown", "facility"]},
            "code":          {"label": "Code",           "required": False, "aliases": ["warehouse code", "id"]},
            "city":          {"label": "City",           "required": False, "aliases": ["location", "place"]},
            "address":       {"label": "Address",        "required": False, "aliases": ["addr"]},
            "manager_name":  {"label": "Manager Name",   "required": False, "aliases": ["manager", "in charge"]},
            "contact_no":    {"label": "Contact No.",    "required": False, "aliases": ["phone", "mobile", "tel"]},
        },
    },
}


def _normalize(s):
    """Normalize a string for fuzzy matching."""
    if not s:
        return ""
    s = str(s).strip().lower()
    s = re.sub(r"[._/\\-]", " ", s)
    s = re.sub(r"\s+", " ", s)
    return s


def _fuzzy_score(a, b):
    """Score between 0-100 for how well string a matches string b."""
    na, nb = _normalize(a), _normalize(b)
    if not na or not nb:
        return 0
    if na == nb:
        return 100
    if na in nb or nb in na:
        return 85
    return int(SequenceMatcher(None, na, nb).ratio() * 100)


def smart_map_columns(headers, entity_type):
    """
    Given a list of raw column headers and an entity type,
    return a mapping dict: { field_name: { col_index, header, confidence, label } }
    plus a list of unmapped columns.
    """
    schema = ENTITY_SCHEMAS.get(entity_type)
    if not schema:
        return {}, list(range(len(headers)))

    fields = schema["fields"]
    mapping = {}
    used_cols = set()

    # For each field, find the best matching column
    for field_name, field_info in fields.items():
        best_score = 0
        best_col = -1
        best_header = ""

        candidates = [field_name.replace("_", " ")] + field_info.get("aliases", []) + [field_info["label"].lower()]

        for col_idx, raw_header in enumerate(headers):
            if col_idx in used_cols:
                continue
            nh = _normalize(raw_header)
            if not nh:
                continue

            for candidate in candidates:
                score = _fuzzy_score(nh, candidate)
                if score > best_score:
                    best_score = score
                    best_col = col_idx
                    best_header = raw_header

        # Only auto-map if confidence is decent
        if best_score >= 55 and best_col >= 0:
            mapping[field_name] = {
                "col_index": best_col,
                "header": best_header,
                "confidence": best_score,
                "label": field_info["label"],
                "required": field_info["required"],
            }
            used_cols.add(best_col)

    unmapped_cols = [
        {"col_index": i, "header": headers[i]}
        for i in range(len(headers))
        if i not in used_cols and _normalize(headers[i])
    ]

    return mapping, unmapped_cols


def detect_entity_type(headers):
    """Guess entity type from column headers. Uses weighted scoring — exact alias matches score higher."""
    scores = {}
    for entity_type, schema in ENTITY_SCHEMAS.items():
        score = 0
        for field_name, field_info in schema["fields"].items():
            candidates = [field_name.replace("_", " ")] + field_info.get("aliases", []) + [field_info["label"].lower()]
            best_match = 0
            for raw_header in headers:
                for cand in candidates:
                    s = _fuzzy_score(raw_header, cand)
                    if s > best_match:
                        best_match = s
            if best_match >= 70:
                score += best_match  # weight by match quality, not just count
        scores[entity_type] = score

    if not scores:
        return "crew"  # default
    best = max(scores, key=scores.get)
    return best if scores[best] > 0 else "crew"


# ──────────────────────────────────────────
# FILE PARSERS — each returns (headers, rows)
# ──────────────────────────────────────────

def parse_xlsx(content):
    """Parse .xlsx/.xls → (headers, rows)."""
    from openpyxl import load_workbook
    wb = load_workbook(filename=BytesIO(content), data_only=True)
    ws = wb[wb.sheetnames[0]]

    # Find the header row (first row with >= 2 non-empty cells)
    header_row = 1
    for r in range(1, min(ws.max_row, 20) + 1):
        vals = [ws.cell(r, c).value for c in range(1, ws.max_column + 1)]
        non_empty = [v for v in vals if v is not None and str(v).strip()]
        if len(non_empty) >= 2:
            header_row = r
            break

    headers = []
    for c in range(1, ws.max_column + 1):
        v = ws.cell(header_row, c).value
        headers.append(str(v).strip() if v else f"Column_{c}")

    rows = []
    for r in range(header_row + 1, ws.max_row + 1):
        row = []
        for c in range(1, ws.max_column + 1):
            v = ws.cell(r, c).value
            row.append(str(v).strip() if v is not None else "")
        # Skip completely empty rows
        if any(cell for cell in row):
            rows.append(row)

    return headers, rows


def parse_csv(content, delimiter=None):
    """Parse .csv/.tsv → (headers, rows)."""
    text = content.decode("utf-8", errors="replace")
    if delimiter is None:
        # Auto-detect delimiter
        sniffer = csv.Sniffer()
        try:
            dialect = sniffer.sniff(text[:2000])
            delimiter = dialect.delimiter
        except Exception:
            delimiter = "," if "," in text[:500] else "\t"

    reader = csv.reader(StringIO(text), delimiter=delimiter)
    all_rows = list(reader)
    if not all_rows:
        return [], []

    headers = [str(h).strip() for h in all_rows[0]]
    rows = []
    for r in all_rows[1:]:
        padded = r + [""] * (len(headers) - len(r))
        if any(cell.strip() for cell in padded):
            rows.append([cell.strip() for cell in padded])

    return headers, rows


def parse_pdf(content):
    """Parse PDF → (headers, rows) by extracting tables."""
    import pdfplumber
    pdf = pdfplumber.open(BytesIO(content))
    all_tables = []
    for page in pdf.pages:
        tables = page.extract_tables()
        for table in tables:
            if table:
                all_tables.extend(table)

    if not all_tables:
        # Fallback: extract lines of text and try to split
        lines = []
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                for line in text.split("\n"):
                    line = line.strip()
                    if line:
                        lines.append(line)
        pdf.close()
        if not lines:
            return [], []
        # Try to parse as pseudo-table (tab or multi-space separated)
        parsed = []
        for line in lines:
            parts = re.split(r"\t|  +", line)
            parts = [p.strip() for p in parts if p.strip()]
            if parts:
                parsed.append(parts)
        if not parsed:
            return [], []
        headers = parsed[0] if parsed else []
        rows = parsed[1:] if len(parsed) > 1 else []
        # Pad rows
        max_cols = max(len(r) for r in ([headers] + rows)) if parsed else 0
        headers += [f"Column_{i+1}" for i in range(len(headers), max_cols)]
        rows = [r + [""] * (max_cols - len(r)) for r in rows]
        return headers, rows

    pdf.close()

    # First non-empty row = headers
    headers = None
    rows = []
    for row in all_tables:
        cleaned = [str(cell).strip() if cell else "" for cell in row]
        if headers is None:
            if any(cleaned):
                headers = cleaned
        else:
            if any(cleaned):
                padded = cleaned + [""] * (len(headers) - len(cleaned))
                rows.append(padded[:len(headers)])

    if not headers:
        return [], []
    return headers, rows


def parse_docx(content):
    """Parse .docx → (headers, rows) by extracting tables."""
    from docx import Document
    doc = Document(BytesIO(content))
    if not doc.tables:
        # Fallback: extract paragraphs as single-column data
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if not lines:
            return [], []
        # Try split by tabs or multiple spaces
        parsed = []
        for line in lines:
            parts = re.split(r"\t|  +", line)
            parts = [p.strip() for p in parts if p.strip()]
            if parts:
                parsed.append(parts)
        if not parsed:
            return [], []
        headers = parsed[0]
        rows = parsed[1:]
        max_cols = max(len(r) for r in ([headers] + rows))
        headers += [f"Column_{i+1}" for i in range(len(headers), max_cols)]
        rows = [r + [""] * (max_cols - len(r)) for r in rows]
        return headers, rows

    # Use the first table
    table = doc.tables[0]
    all_rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        all_rows.append(cells)

    if not all_rows:
        return [], []

    headers = all_rows[0]
    rows = [r for r in all_rows[1:] if any(r)]
    return headers, rows


def parse_image(content):
    """Parse image via OCR → (headers, rows)."""
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        return [], []

    img = Image.open(BytesIO(content))

    # Try to get structured text
    text = pytesseract.image_to_string(img)
    if not text.strip():
        return [], []

    lines = [l.strip() for l in text.split("\n") if l.strip()]
    if not lines:
        return [], []

    # Try to split into columns by tabs or multiple spaces
    parsed = []
    for line in lines:
        parts = re.split(r"\t|  +|\|", line)
        parts = [p.strip() for p in parts if p.strip()]
        if parts:
            parsed.append(parts)

    if not parsed:
        return [], []

    # First row with >= 2 cells = header
    header_idx = 0
    for i, row in enumerate(parsed):
        if len(row) >= 2:
            header_idx = i
            break

    headers = parsed[header_idx]
    rows = parsed[header_idx + 1:]
    max_cols = max(len(r) for r in ([headers] + rows)) if parsed else 0
    headers += [f"Column_{i+1}" for i in range(len(headers), max_cols)]
    rows = [r + [""] * (max_cols - len(r)) for r in rows]
    return headers, rows


# ──────────────────────────────────────────
# MAIN ENTRY POINT
# ──────────────────────────────────────────

SUPPORTED_EXTENSIONS = {
    ".xlsx": parse_xlsx,
    ".xls": parse_xlsx,
    ".csv": parse_csv,
    ".tsv": lambda c: parse_csv(c, delimiter="\t"),
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".doc": parse_docx,
    ".jpg": parse_image,
    ".jpeg": parse_image,
    ".png": parse_image,
    ".bmp": parse_image,
    ".tiff": parse_image,
    ".tif": parse_image,
}


def parse_any_file(filename, content):
    """
    Parse any supported file and return (headers, rows).
    Raises ValueError if format is unsupported.
    """
    import os
    ext = os.path.splitext(filename)[1].lower()
    parser = SUPPORTED_EXTENSIONS.get(ext)
    if not parser:
        raise ValueError(f"Unsupported file format: {ext}. Supported: {', '.join(SUPPORTED_EXTENSIONS.keys())}")
    return parser(content)


def smart_preview(filename, content, entity_type=None):
    """
    Full smart preview: parse file → detect entity → map columns → return preview.

    Returns dict:
    {
        "headers": [...],
        "detected_entity_type": "crew",
        "mapping": { field_name: { col_index, header, confidence, label, required } },
        "unmapped_columns": [ { col_index, header } ],
        "available_fields": { field_name: { label, required } },
        "sample_rows": [[...], ...],  # first 10 rows
        "total_rows": int,
    }
    """
    headers, rows = parse_any_file(filename, content)
    if not headers:
        raise ValueError("Could not detect any column headers in the file.")

    # Auto-detect entity type if not provided
    if not entity_type:
        entity_type = detect_entity_type(headers)

    mapping, unmapped = smart_map_columns(headers, entity_type)

    # Available fields for user to pick from (for unmapped cols)
    schema = ENTITY_SCHEMAS.get(entity_type, {})
    mapped_fields = set(mapping.keys())
    available_fields = {}
    for fname, finfo in schema.get("fields", {}).items():
        available_fields[fname] = {
            "label": finfo["label"],
            "required": finfo["required"],
            "mapped": fname in mapped_fields,
        }

    return {
        "headers": headers,
        "detected_entity_type": entity_type,
        "mapping": mapping,
        "unmapped_columns": unmapped,
        "available_fields": available_fields,
        "sample_rows": rows[:10],
        "total_rows": len(rows),
    }


def apply_mapping_and_import(filename, content, entity_type, user_mapping, db):
    """
    Apply user-confirmed column mapping and import rows.

    user_mapping: { field_name: col_index } — final mapping from user
    Returns: { count, accepted, rejected }
    """
    from . import models
    from .codegen import next_employee_code, next_vendor_code, next_client_code, next_asset_code

    headers, rows = parse_any_file(filename, content)
    accepted = []
    rejected = []
    count = 0

    def _get(row, field_name):
        col_idx = user_mapping.get(field_name)
        if col_idx is None or col_idx < 0 or col_idx >= len(row):
            return ""
        return str(row[col_idx]).strip() if row[col_idx] else ""

    for row_idx, row in enumerate(rows, start=1):
        try:
            with db.begin_nested():
                if entity_type == "crew":
                    name = _get(row, "full_name")
                    if not name:
                        rejected.append({"row": row_idx, "reason": "Missing name"})
                        continue
                    code = _get(row, "employee_code") or next_employee_code(db)
                    if db.query(models.CrewMember).filter(models.CrewMember.employee_code == code).first():
                        rejected.append({"row": row_idx, "reason": f"Duplicate: {code}"})
                        continue
                    db.add(models.CrewMember(
                        employee_code=code, full_name=name,
                        role=_get(row, "role") or "Crew",
                        manpower_type=_get(row, "manpower_type") or "inhouse",
                        home_station=_get(row, "home_station") or "Base",
                        phone=_get(row, "phone") or None,
                        address=_get(row, "address") or None,
                        aadhar_number=_get(row, "aadhar_number") or None,
                        id_proof_type=_get(row, "id_proof_type") or None,
                        id_proof_number=_get(row, "id_proof_number") or None,
                        status="available",
                    ))
                    db.flush()
                    accepted.append({"row": row_idx, "name": name, "code": code})
                    count += 1

                elif entity_type == "clients":
                    name = _get(row, "name")
                    if not name:
                        rejected.append({"row": row_idx, "reason": "Missing name"})
                        continue
                    code = _get(row, "client_code") or next_client_code(db)
                    if db.query(models.Client).filter(models.Client.client_code == code).first():
                        rejected.append({"row": row_idx, "reason": f"Duplicate: {code}"})
                        continue
                    db.add(models.Client(
                        client_code=code, name=name,
                        industry_type=_get(row, "industry_type") or None,
                        billing_address=_get(row, "billing_address") or None,
                        gst_number=_get(row, "gst_number") or None,
                        notes=_get(row, "notes") or None,
                    ))
                    db.flush()
                    accepted.append({"row": row_idx, "name": name, "code": code})
                    count += 1

                elif entity_type == "vendors":
                    name = _get(row, "name")
                    if not name:
                        rejected.append({"row": row_idx, "reason": "Missing name"})
                        continue
                    code = _get(row, "vendor_code") or next_vendor_code(db)
                    if db.query(models.Vendor).filter(models.Vendor.vendor_code == code).first():
                        rejected.append({"row": row_idx, "reason": f"Duplicate: {code}"})
                        continue
                    db.add(models.Vendor(
                        vendor_code=code, name=name,
                        vendor_type=_get(row, "vendor_type") or "equipment",
                        city=_get(row, "city") or None,
                        contact_person=_get(row, "contact_person") or None,
                        phone=_get(row, "phone") or None,
                        email=_get(row, "email") or None,
                        gst_number=_get(row, "gst_number") or None,
                        notes=_get(row, "notes") or None,
                    ))
                    db.flush()
                    accepted.append({"row": row_idx, "name": name, "code": code})
                    count += 1

                elif entity_type == "warehouses":
                    name = _get(row, "name")
                    if not name:
                        rejected.append({"row": row_idx, "reason": "Missing name"})
                        continue
                    code = _get(row, "code") or f"WH-{row_idx}"
                    if db.query(models.Warehouse).filter(models.Warehouse.code == code).first():
                        rejected.append({"row": row_idx, "reason": f"Duplicate: {code}"})
                        continue
                    db.add(models.Warehouse(
                        code=code, name=name,
                        city=_get(row, "city") or "City",
                        address=_get(row, "address") or None,
                        manager_name=_get(row, "manager_name") or None,
                        contact_no=_get(row, "contact_no") or None,
                    ))
                    db.flush()
                    accepted.append({"row": row_idx, "name": name, "code": code})
                    count += 1

                elif entity_type == "inventory":
                    name = _get(row, "name")
                    if not name:
                        rejected.append({"row": row_idx, "reason": "Missing name"})
                        continue
                    code = _get(row, "asset_code") or next_asset_code(db)
                    if db.query(models.InventoryItem).filter(models.InventoryItem.asset_code == code).first():
                        rejected.append({"row": row_idx, "reason": f"Duplicate: {code}"})
                        continue
                    db.add(models.InventoryItem(
                        asset_code=code, name=name,
                        serial_number=_get(row, "serial_number") or None,
                        category=_get(row, "category") or None,
                        item_type=_get(row, "item_type") or "device",
                        brand=_get(row, "brand") or None,
                        model_no=_get(row, "model_no") or None,
                        location_text=_get(row, "location_text") or None,
                        warranty_expiry=_get(row, "warranty_expiry") or None,
                        purchase_date=_get(row, "purchase_date") or None,
                        notes=_get(row, "notes") or None,
                        status="available",
                    ))
                    db.flush()
                    accepted.append({"row": row_idx, "name": name, "code": code})
                    count += 1

        except Exception as e:
            rejected.append({"row": row_idx, "reason": str(e)})

    db.commit()
    return {"count": count, "accepted": accepted, "rejected": rejected}
